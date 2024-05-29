import docx
from docx.enum.style import WD_STYLE_TYPE

guests = open('guests.txt', 'r')
guest_array = guests.readlines()

doc = docx.Document('invitation.docx')
doc._body.clear_content()

for i, guest in enumerate(guest_array):
    paragraph1 = doc.add_paragraph('It would be a pleasure to have the company of')
    paragraph1.style = doc.styles['invitation1']
    paragraph2 = doc.add_paragraph(guest[:-1])
    paragraph2.style = doc.styles['invitation2']
    paragraph3 = doc.add_paragraph('at 11010 Memory Lane on the Evening of')
    paragraph3.style = doc.styles['invitation1']
    paragraph4 = doc.add_paragraph('April 1st')
    paragraph4.style = doc.styles['invitation3']
    paragraph5 = doc.add_paragraph('at 7 o\'clock')
    paragraph5.style = doc.styles['invitation1']
    if i != (len(guest_array) - 1):
        paragraph6 = doc.add_page_break()

doc.save('invitation.docx')